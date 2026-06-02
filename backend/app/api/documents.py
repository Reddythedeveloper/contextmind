from fastapi import APIRouter, UploadFile, File, Depends, HTTPException
from app.services.rag_service import RAGService
from app.services.embedding_service import EmbeddingService
from app.services.vector_store import VectorStore
import PyPDF2
import docx
import io
import uuid

from app.models.document import Document
from app.models.conversation import Conversation
from sqlalchemy.future import select
from sqlalchemy.ext.asyncio import AsyncSession
from app.core.database import get_db

router = APIRouter(prefix="/documents", tags=["documents"])

# Dependency injection
def get_rag_service():
    embedder = EmbeddingService()
    store = VectorStore()
    return RAGService(embedder, store)

# Mock user dependency
async def get_current_user():
    return {"id": "user-123", "email": "test@example.com"}

@router.post("/upload")
async def upload_document(
    session_id: str,
    file: UploadFile = File(...),
    user: dict = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    rag_service: RAGService = Depends(get_rag_service)
):
    # Find conversation
    result = await db.execute(select(Conversation).where(Conversation.session_id == session_id))
    conv = result.scalar_one_or_none()
    if not conv:
        conv = Conversation(session_id=session_id, user_id=user["id"])
        db.add(conv)
        await db.commit()
        await db.refresh(conv)

    content = await file.read()
    filename = file.filename
    text = ""
    # ... extraction logic ...
    if filename.endswith(".pdf"):
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        for page in pdf_reader.pages:
            text += page.extract_text()
    elif filename.endswith(".docx"):
        doc = docx.Document(io.BytesIO(content))
        for para in doc.paragraphs:
            text += para.text + "\n"
    elif filename.endswith(".txt"):
        text = content.decode("utf-8")
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    doc_id = str(uuid.uuid4())
    num_chunks = await rag_service.ingest(
        text=text,
        doc_id=doc_id,
        user_id=user["id"],
        session_id=session_id,
        metadata={"filename": filename}
    )

    # Save document metadata linked to conversation
    new_doc = Document(
        conversation_id=conv.id,
        user_id=user["id"],
        filename=filename,
        doc_id=doc_id
    )
    db.add(new_doc)
    await db.commit()

    return {
        "message": "Document uploaded and indexed",
        "doc_id": doc_id,
        "chunks": num_chunks
    }

@router.get("/")
async def get_documents(
    session_id: str,
    db: AsyncSession = Depends(get_db)
):
    result = await db.execute(
        select(Document)
        .join(Conversation)
        .where(Conversation.session_id == session_id)
    )
    return result.scalars().all()

@router.get("/search")
async def search_documents(
    q: str,
    user: dict = Depends(get_current_user),
    rag_service: RAGService = Depends(get_rag_service)
):
    results = await rag_service.retrieve(query=q, user_id=user["id"])
    return {"results": results}
